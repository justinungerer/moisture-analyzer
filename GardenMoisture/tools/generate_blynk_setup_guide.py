from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Blynk_Setup_Guide_Garden_Moisture.docx"


DATASTREAMS = []
for zone in range(1, 16):
    DATASTREAMS.append(
        (
            f"Zone {zone:02d} Moisture",
            f"V{zone - 1}",
            "Integer",
            "0",
            "100",
            "0",
            "%",
            "Off",
            "Gauge / chart telemetry",
        )
    )

DATASTREAMS.extend(
    [
        ("Battery Level", "V15", "Integer", "0", "100", "0", "%", "Off", "Telemetry"),
        ("Sleep Interval", "V16", "Integer", "5", "1440", "30", "min", "On", "Global setting"),
        ("Moisture Alert Threshold", "V17", "Integer", "0", "100", "25", "%", "On", "Global setting"),
        ("Selected Zone Dry Raw", "V18", "Integer", "0", "4095", "3200", "ADC", "Off", "Per-zone live edit"),
        ("Selected Zone Wet Raw", "V19", "Integer", "0", "4095", "1400", "ADC", "Off", "Per-zone live edit"),
        ("Calibration Mode", "V20", "Integer", "0", "1", "0", "", "On", "Global setting"),
        ("Selected Zone", "V21", "Integer", "1", "15", "1", "", "On", "Global setting"),
        ("Capture Dry", "V22", "Integer", "0", "1", "0", "", "Off", "Momentary command"),
        ("Capture Wet", "V23", "Integer", "0", "1", "0", "", "Off", "Momentary command"),
        ("Selected Zone Live Raw", "V24", "Integer", "0", "4095", "0", "ADC", "Off", "Telemetry"),
        ("Stay Awake", "V25", "Integer", "0", "1", "0", "", "On", "Global setting"),
        ("Report Now", "V26", "Integer", "0", "1", "0", "", "Off", "Momentary command"),
        ("Selected Zone Enabled", "V27", "Integer", "0", "1", "1", "", "Off", "Per-zone live edit"),
        ("Battery Voltage", "V28", "Double", "0", "5", "0", "V", "Off", "Telemetry; show 2 decimals"),
        ("Low-Battery Threshold", "V29", "Integer", "0", "100", "15", "%", "On", "Global setting"),
        ("Diagnostics", "V30", "String", "—", "—", "", "", "Off", "Telemetry text"),
    ]
)


WEB_WIDGETS = [
    ("Overview", "Gauge", "Battery Level", "V15", "0–100%; title “Battery”"),
    ("Overview", "Label / Value Display", "Battery Voltage", "V28", "Show 2 decimal places and V"),
    ("Overview", "Numeric Input", "Sleep Interval", "V16", "Min 5, max 1440, step 5"),
    ("Overview", "Numeric Input", "Moisture Alert Threshold", "V17", "Min 0, max 100, step 1"),
    ("Overview", "Numeric Input", "Low-Battery Threshold", "V29", "Min 0, max 100, step 1"),
    ("Overview", "Switch", "Stay Awake", "V25", "OFF=0, ON=1"),
    ("Overview", "Button", "Report Now", "V26", "Mode Push/Momentary; OFF=0, ON=1"),
    ("Overview", "Label", "Diagnostics", "V30", "Wide widget; allow text wrapping"),
    ("Moisture", "15 Gauges", "Zone 01–15 Moisture", "V0–V14", "One gauge per zone; 0–100%"),
    ("Moisture", "Chart (optional)", "Zone moisture history", "V0–V14", "Add selected zones; 15 lines may be crowded"),
    ("Calibration", "Switch", "Calibration Mode", "V20", "OFF=0, ON=1"),
    ("Calibration", "Numeric Input", "Selected Zone", "V21", "Min 1, max 15, step 1"),
    ("Calibration", "Switch", "Selected Zone Enabled", "V27", "OFF=0, ON=1"),
    ("Calibration", "Label / Value Display", "Selected Zone Live Raw", "V24", "0–4095 ADC"),
    ("Calibration", "Numeric Input", "Selected Zone Dry Raw", "V18", "Min 0, max 4095, step 1"),
    ("Calibration", "Numeric Input", "Selected Zone Wet Raw", "V19", "Min 0, max 4095, step 1"),
    ("Calibration", "Button", "Capture Dry", "V22", "Mode Push/Momentary; OFF=0, ON=1"),
    ("Calibration", "Button", "Capture Wet", "V23", "Mode Push/Momentary; OFF=0, ON=1"),
]


MOBILE_WIDGETS = [
    ("Header", "Battery Level", "Battery Level", "V15", "Optional header widget; 0–100%"),
    ("Overview", "Gauge", "Battery Level", "V15", "Title “Battery”"),
    ("Overview", "Labeled Value", "Battery Voltage", "V28", "2 decimal places"),
    ("Overview", "Numeric Input", "Sleep Interval", "V16", "5–1440, step 5"),
    ("Overview", "Numeric Input", "Moisture Alert Threshold", "V17", "0–100, step 1"),
    ("Overview", "Numeric Input", "Low-Battery Threshold", "V29", "0–100, step 1"),
    ("Overview", "Switch", "Stay Awake", "V25", "OFF=0, ON=1"),
    ("Overview", "Button", "Report Now", "V26", "Mode Push; OFF=0, ON=1"),
    ("Overview", "Labeled Value", "Diagnostics", "V30", "Use a wide widget"),
    ("Zones 1–8", "Gauge", "Zone 01–08 Moisture", "V0–V7", "One gauge per zone"),
    ("Zones 9–15", "Gauge", "Zone 09–15 Moisture", "V8–V14", "One gauge per zone"),
    ("History (optional)", "SuperChart", "Moisture history", "V0–V14", "Use only key zones if the chart is crowded"),
    ("Calibration", "Switch", "Calibration Mode", "V20", "OFF=0, ON=1"),
    ("Calibration", "Numeric Input", "Selected Zone", "V21", "1–15, step 1"),
    ("Calibration", "Switch", "Selected Zone Enabled", "V27", "OFF=0, ON=1"),
    ("Calibration", "Labeled Value", "Selected Zone Live Raw", "V24", "0–4095 ADC"),
    ("Calibration", "Numeric Input", "Selected Zone Dry Raw", "V18", "0–4095, step 1"),
    ("Calibration", "Numeric Input", "Selected Zone Wet Raw", "V19", "0–4095, step 1"),
    ("Calibration", "Button", "Capture Dry", "V22", "Mode Push; OFF=0, ON=1"),
    ("Calibration", "Button", "Capture Wet", "V23", "Mode Push; OFF=0, ON=1"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, header_fill="2F75B5"):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        set_cell_text_color(cell, "FFFFFF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_index, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cell, "EAF2F8")


def add_numbered_steps(document, steps, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    for step in steps:
        document.add_paragraph(step, style=style)


def add_bullets(document, bullets, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for bullet in bullets:
        document.add_paragraph(bullet, style=style)


def add_note(document, title, text, fill="FFF2CC"):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{title}: ")
    run.bold = True
    paragraph.add_run(text)
    document.add_paragraph()


def add_widget_table(document, rows):
    headers = ["Page / area", "Widget", "Datastream", "Pin", "Required settings"]
    table = document.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(table, "548235")
    return table


document = Document()
section = document.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = document.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(24)
styles["Title"].font.color.rgb = RGBColor(31, 78, 121)
for style_name, size, color in [
    ("Heading 1", 18, "1F4E79"),
    ("Heading 2", 14, "2F75B5"),
    ("Heading 3", 11, "548235"),
]:
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.size = Pt(size)
    styles[style_name].font.color.rgb = RGBColor.from_string(color)

title = document.add_paragraph()
title.style = "Title"
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Blynk Setup Guide\nGarden Moisture Monitor")
subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Click-by-click instructions for Blynk.Console and the Blynk IoT mobile app")
run.italic = True
run.font.size = Pt(12)

summary = document.add_table(rows=4, cols=2)
summary.style = "Light Shading Accent 1"
summary.cell(0, 0).text = "Blynk template name"
summary.cell(0, 1).text = "Soil Moisture Monitoring"
summary.cell(1, 0).text = "Template ID in firmware"
summary.cell(1, 1).text = "TMPL2c6XdcWxv"
summary.cell(2, 0).text = "Firmware"
summary.cell(2, 1).text = "Garden Moisture Monitor v1.3.0"
summary.cell(3, 0).text = "Virtual pins"
summary.cell(3, 1).text = "V0 through V30"
document.add_paragraph()

add_note(
    document,
    "Important before changing anything",
    "Connectivity already works. Use the existing template and device. Do not create a replacement template, "
    "change the Template ID, delete the device, or regenerate its authentication token. In Blynk.Console, verify "
    "that the existing template ID is TMPL2c6XdcWxv.",
    "FCE4D6",
)
add_note(
    document,
    "First boot after flashing v1.3.0",
    "Let the device complete one successful Blynk connection before changing settings. On that first connection "
    "the firmware seeds the dashboard with safe defaults once. Later connections pull the latest global settings "
    "from Blynk instead of overwriting them.",
)

intro = document.add_paragraph()
intro.add_run("How Blynk is organized. ").bold = True
intro.add_run(
    "Datastreams and events belong to a Device Template and are created in the web-based Blynk.Console. "
    "The Web Dashboard and Mobile Dashboard are independent layouts, but both use the same datastreams. "
    "Therefore, complete Section 1 first, even if you mainly intend to use the phone app."
)

# SECTION 1
document.add_heading("Section 1 — Web app (Blynk.Console)", level=1)

document.add_heading("1. Sign in and verify the existing template", level=2)
add_numbered_steps(
    document,
    [
        "Open a web browser and go to https://blynk.cloud.",
        "Sign in with the same Blynk account that currently owns the moisture analyzer.",
        "Open the user/profile menu and turn Developer Mode ON if it is not already enabled.",
        "In the left navigation, click Developer Zone, then click My Templates (some accounts show Templates).",
        "Click the existing template named Soil Moisture Monitoring.",
        "Open its Info or Firmware Configuration area and verify Template ID = TMPL2c6XdcWxv.",
        "If that ID matches, continue. If it does not match, stop: you are editing the wrong template.",
    ],
)

document.add_heading("2. Create or audit the shared datastreams", level=2)
add_note(
    document,
    "Do not make duplicates",
    "If a virtual pin already exists, click it and correct its settings instead of creating another datastream "
    "on the same pin. Each virtual pin V0–V30 should appear only once.",
)
add_numbered_steps(
    document,
    [
        "Inside the Soil Moisture Monitoring template, click the Datastreams tab.",
        "Compare the existing list against the complete table below.",
        "For each missing row, click + New Datastream (or New Datastream), then choose Virtual Pin.",
        "Enter the Name exactly as shown, choose the listed Virtual Pin, choose the Data Type, and enter Min, Max, Default, and Units.",
        "Open Advanced Settings. Turn Sync with latest server value ON only where the table says On; leave it OFF elsewhere.",
        "Leave Invalidate Value OFF unless you specifically want a widget to become blank when data is old.",
        "Click Create or Save.",
        "Repeat until every pin V0 through V30 exists exactly once.",
    ],
)

headers = ["Datastream name", "Pin", "Type", "Min", "Max", "Default", "Unit", "Sync latest", "Use"]
landscape = document.add_section(WD_SECTION.NEW_PAGE)
landscape.orientation = WD_ORIENT.LANDSCAPE
landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
landscape.top_margin = Inches(0.55)
landscape.bottom_margin = Inches(0.55)
landscape.left_margin = Inches(0.45)
landscape.right_margin = Inches(0.45)
table = document.add_table(rows=1, cols=len(headers))
for index, value in enumerate(headers):
    table.rows[0].cells[index].text = value
for row in DATASTREAMS:
    cells = table.add_row().cells
    for index, value in enumerate(row):
        cells[index].text = str(value)
style_table(table)

portrait = document.add_section(WD_SECTION.NEW_PAGE)
portrait.orientation = WD_ORIENT.PORTRAIT
portrait.page_width, portrait.page_height = portrait.page_height, portrait.page_width
portrait.top_margin = Inches(0.65)
portrait.bottom_margin = Inches(0.65)
portrait.left_margin = Inches(0.65)
portrait.right_margin = Inches(0.65)

document.add_heading("3. Create the two firmware events and notifications", level=2)
add_numbered_steps(
    document,
    [
        "Still inside the template, click the Events tab.",
        "Click + Create Event or New Event.",
        "For the first event, enter Name = Low Moisture and Event Code = low_moisture.",
        "Choose Warning (or the closest warning-level event type) and save the event.",
        "Open the event again and click its Notifications tab.",
        "Enable Notifications, enable Push, and select the device Owner/current user as the recipient. Enable Email too if your plan and account provide it.",
        "Leave notification frequency limits disabled: the firmware already limits moisture alerts once per zone per local day, and a template-wide one-per-day limit would incorrectly hide alerts from other zones.",
        "Return to the Events list and create Name = Low Battery with Event Code = low_battery.",
        "Choose Warning, save it, then enable Push notifications for the Owner/current user.",
        "Leave notification frequency limits disabled for Low Battery as well; the firmware already rate-limits it.",
    ],
)
add_note(
    document,
    "Event codes are exact",
    "Use lowercase with the underscore exactly as shown: low_moisture and low_battery. The firmware calls those "
    "exact strings. Changing capitalization or spaces will prevent alerts.",
    "FCE4D6",
)

document.add_heading("4. Build the Web Dashboard", level=2)
add_numbered_steps(
    document,
    [
        "Inside the same template, click the Web Dashboard tab.",
        "Click Edit in the upper-right corner to open the dashboard editor and Widget Box.",
        "Create three dashboard tabs/pages if your Blynk plan shows tabs: Overview, Moisture, and Calibration. If tabs are unavailable, arrange the same groups vertically on one dashboard.",
        "For each row in the widget plan below, drag the specified widget from the Widget Box to the named page.",
        "Hover over the new widget and click its gear icon.",
        "Choose the listed Datastream, set the title, apply the required settings, and click Save/Done.",
        "Resize and arrange widgets after they are configured.",
    ],
)
add_widget_table(document, WEB_WIDGETS)

document.add_heading("5. Add all 15 web moisture gauges", level=2)
add_numbered_steps(
    document,
    [
        "Open the Moisture dashboard page while still in Edit mode.",
        "Drag a Gauge widget onto the canvas.",
        "Hover over it, click the gear icon, and select Zone 01 Moisture (V0).",
        "Set the title to Zone 1, keep the datastream range 0–100 and unit %, then save.",
        "Repeat for Zone 02 Moisture (V1) through Zone 15 Moisture (V14).",
        "Arrange the gauges in a readable grid, such as five columns by three rows.",
        "Optionally add a multi-datastream Chart. Add only the zones whose history matters most if 15 lines are too crowded.",
    ],
)

document.add_heading("6. Save, apply, and view live web data", level=2)
add_numbered_steps(
    document,
    [
        "Click Save and Apply in the Web Dashboard/template editor. If Blynk shows separate Save and Apply buttons, click both.",
        "Leave Developer Zone and use Search or My Devices in the main navigation.",
        "Open the existing garden moisture device.",
        "Open its Dashboard tab. This device view—not the template preview—is where actual live values appear.",
        "If the device is sleeping, widgets may show the last reported values until its next scheduled wake.",
        "To keep it online, turn Stay Awake ON. The firmware will receive that stored value at its next connection and remain awake.",
    ],
)

document.add_heading("7. Web dashboard test", level=2)
add_bullets(
    document,
    [
        "Battery Level shows 0–100% on V15 and Battery Voltage shows a value near the actual single-cell voltage on V28.",
        "All 15 moisture gauges show percentages, never raw ADC values.",
        "Change Sleep Interval to 35; after the device connects, refresh the page and confirm it remains 35.",
        "Press Report Now once. The button should return to OFF/0 automatically and telemetry should refresh.",
        "Turn Stay Awake ON before doing calibration; turn it OFF when finished to restore battery-saving sleep.",
        "Diagnostics should show uptime, RSSI, free heap, battery, wake reason, and firmware version.",
    ],
)

# SECTION 2
document.add_page_break()
document.add_heading("Section 2 — Mobile app (Blynk IoT)", level=1)

document.add_heading("1. Install, sign in, and enable Developer Mode", level=2)
add_numbered_steps(
    document,
    [
        "Install Blynk IoT from Google Play on Android or the App Store on iPhone/iPad.",
        "Open Blynk IoT and sign in with the same account used at https://blynk.cloud.",
        "Open Profile or the account menu.",
        "Turn Developer Mode ON. The exact location can vary slightly by app version, but it is under the profile/account area.",
        "Return to the main screen. Open Developer Mode and tap the Soil Moisture Monitoring template. Do not create a new template or device.",
        "You are now editing the Mobile Dashboard for that template. This layout is separate from the Web Dashboard.",
    ],
)

document.add_heading("2. Understand the mobile editor", level=2)
add_bullets(
    document,
    [
        "Tap an empty area of the dashboard or tap the + icon in the upper-right to open the Widget Box.",
        "Tap a widget in the Widget Box to place it in the first available space.",
        "Tap a placed widget to open its settings, then choose its Datastream and other options.",
        "Drag a widget to move it. Use the resize handles or long-press/resize gesture offered by your app version.",
        "Mobile widgets use the datastreams created in Section 1; datastreams themselves cannot be fully configured in the mobile dashboard editor.",
    ],
)

document.add_heading("3. Create the recommended mobile layout", level=2)
add_numbered_steps(
    document,
    [
        "Add a Tabs widget if it is available in your plan/app version.",
        "Create pages named Overview, Zones 1–8, Zones 9–15, and Calibration. Add History as an optional fifth page.",
        "If Tabs is unavailable, place the same widget groups vertically and use Text widgets as section headings.",
        "Add and configure the widgets in the table below one at a time.",
    ],
)
add_widget_table(document, MOBILE_WIDGETS)

document.add_heading("4. Configure a typical mobile gauge", level=2)
add_numbered_steps(
    document,
    [
        "Open the Zones 1–8 page.",
        "Tap +, find Gauge (or Enhanced/Radial Gauge if Gauge is not offered), and tap it.",
        "Tap the new gauge to open Widget Settings.",
        "Set Title = Zone 1.",
        "Tap Datastream and select Zone 01 Moisture (V0).",
        "Use the datastream's default range 0–100 and unit %. Optionally choose a green-to-red ramp where higher moisture is green.",
        "Tap Back/Done to save the widget.",
        "Repeat for Zones 2–8 using V1–V7, then repeat on Zones 9–15 using V8–V14.",
    ],
)

document.add_heading("5. Configure mobile buttons and switches correctly", level=2)
add_numbered_steps(
    document,
    [
        "For Stay Awake (V25), add a Switch widget. Set OFF value = 0 and ON value = 1.",
        "For Calibration Mode (V20), add a Switch widget. Set OFF = 0 and ON = 1.",
        "For Selected Zone Enabled (V27), add a Switch widget. Set OFF = 0 and ON = 1.",
        "For Report Now (V26), add a Button widget and choose Push/Momentary mode—not Switch mode. Set OFF = 0 and ON = 1.",
        "For Capture Dry (V22) and Capture Wet (V23), use Button widgets in Push/Momentary mode with OFF = 0 and ON = 1.",
        "The firmware writes the three momentary buttons back to 0 after handling them; seeing them release is expected.",
    ],
)

document.add_heading("6. Configure numeric inputs and displays", level=2)
add_numbered_steps(
    document,
    [
        "Add Numeric Input for Sleep Interval (V16): minimum 5, maximum 1440, step 5.",
        "Add Numeric Input for Moisture Alert Threshold (V17): minimum 0, maximum 100, step 1.",
        "Add Numeric Input for Low-Battery Threshold (V29): minimum 0, maximum 100, step 1.",
        "Add Numeric Input for Selected Zone (V21): minimum 1, maximum 15, step 1.",
        "Add Numeric Input for Dry Raw (V18) and Wet Raw (V19): minimum 0, maximum 4095, step 1.",
        "Add Labeled Value for Live Raw (V24), Battery Voltage (V28), and Diagnostics (V30).",
        "If your app does not offer Labeled Value for the String datastream V30, use Value Display or another text-capable display widget.",
    ],
)

document.add_heading("7. Save and open the actual mobile device dashboard", level=2)
add_numbered_steps(
    document,
    [
        "After adding widgets, exit the editor with Back, Done, or the editor's exit icon. Blynk normally saves the mobile layout as you leave.",
        "Turn Developer Mode OFF if you want to prevent accidental layout changes during everyday use.",
        "Return to the device list and tap the existing garden moisture device tile.",
        "Confirm the configured dashboard opens and shows the same datastream values as the web dashboard.",
        "If values look stale, check the device's online/offline status. Deep sleep intentionally makes it offline between reports.",
    ],
)

document.add_heading("8. Safe calibration procedure from either dashboard", level=2)
add_numbered_steps(
    document,
    [
        "Turn Stay Awake (V25) ON. If the device is asleep, wait for its next scheduled wake; it will sync this stored setting and remain online.",
        "Wait until the device status shows Online.",
        "Turn Calibration Mode (V20) ON.",
        "Set Selected Zone (V21) to the zone number you want to calibrate.",
        "Check Selected Zone Enabled (V27). Leave it ON for an active sensor; turn it OFF only for unused/broken zones that should not alert.",
        "Expose that zone's sensor to its dry reference condition and wait for Live Raw (V24) to stabilize.",
        "Tap Capture Dry (V22) once. Confirm Selected Zone Dry Raw (V18) updates.",
        "Place the sensor in its wet reference condition, wait for Live Raw (V24) to stabilize, then tap Capture Wet (V23) once.",
        "Confirm Selected Zone Wet Raw (V19) updates. Capacitive sensors commonly read lower when wet; the firmware also supports inverted readings.",
        "Repeat for each zone.",
        "Turn Calibration Mode OFF, then turn Stay Awake OFF. The normal awake window will run and the device will return to deep sleep.",
    ],
)

document.add_heading("9. Final end-to-end checklist", level=2)
checklist = [
    "□ Existing template ID is TMPL2c6XdcWxv; connectivity credentials were not changed.",
    "□ Every pin V0–V30 exists once, with the correct type and range.",
    "□ Sync latest is ON only for V16, V17, V20, V21, V25, and V29.",
    "□ Event codes are exactly low_moisture and low_battery, and Push is enabled.",
    "□ V22, V23, and V26 use Push/Momentary buttons, not latching switches.",
    "□ V20, V25, and V27 use switches with OFF=0 and ON=1.",
    "□ All V0–V14 widgets show percentages.",
    "□ Web Dashboard changes were saved/applied and live data is viewed from My Devices.",
    "□ Mobile dashboard was built separately and opened from the actual device tile.",
    "□ Stay Awake and Calibration Mode are OFF after setup to protect the battery.",
]
for item in checklist:
    document.add_paragraph(item)

add_note(
    document,
    "If a widget is unavailable",
    "Blynk widget availability and dashboard capacity depend on the account plan and app version. Substitute a "
    "Gauge with Labeled Value, or reduce optional charts first. Do not change the datastream pin or data type.",
)

document.add_heading("10. Troubleshooting", level=2)
troubleshooting = [
    ("A control snaps back to its old value", "Verify Sync latest is ON for V16, V17, V20, V21, V25, or V29, then allow the device to connect. Also confirm you are using the existing template."),
    ("A button stays ON", "Edit V22, V23, or V26 and select Push/Momentary mode with 0/1 values."),
    ("No live values in template editor", "Open Search/My Devices → the actual device → Dashboard. Template editors are layout previews."),
    ("Device is often Offline", "That is normal during deep sleep. Use Stay Awake for setup and switch it OFF afterward."),
    ("Calibration controls do nothing", "Turn Stay Awake ON, wait for Online, then turn Calibration Mode ON."),
    ("No alert notification", "Verify the exact event code, enable Notifications/Push, select a recipient, and allow phone notifications for Blynk IoT in the phone's operating-system settings."),
    ("Battery voltage is clearly wrong", "Confirm the physical divider ratio and single-cell battery configuration before changing firmware constants."),
]
table = document.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Symptom"
table.rows[0].cells[1].text = "What to check"
for symptom, fix in troubleshooting:
    cells = table.add_row().cells
    cells[0].text = symptom
    cells[1].text = fix
style_table(table, "C65911")

document.add_paragraph()
source = document.add_paragraph()
source.add_run("Prepared for the Garden Moisture Monitor v1.3.0 firmware. ").bold = True
source.add_run(
    "Blynk menu wording can vary slightly by account plan, Android/iOS version, and future interface updates. "
    "The guide follows current Blynk documentation terminology: Developer Zone → My Templates → Datastreams / "
    "Events / Web Dashboard, and Developer Mode in the Blynk IoT mobile app."
)

document.save(OUTPUT)
print(OUTPUT)
